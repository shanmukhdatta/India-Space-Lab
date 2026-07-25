import os
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = r"c:\Users\shanm\ISL\India-Space-Lab-main\Rocketry_Project\Rocketry_FEM_CFD_Project_Report.docx"
OUT_DIR = r"c:\Users\shanm\ISL\India-Space-Lab-main\Rocketry_Project"

doc = docx.Document(DOC_PATH)

for i, p in enumerate(doc.paragraphs):
    if "PLACEHOLDER" in p.text:
        print(f"Found placeholder at paragraph {i}: {p.text[:60]}...")
        if "(1) Von Mises stress contour" in p.text:
            p.text = "SimScale Finite Element Analysis (FEA) Results — Von Mises Bending Stress & Displacement Contours:"
            p.runs[0].font.bold = True
            
            # Add images
            p1 = p.insert_paragraph_before()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p1.add_run()
            run1.add_picture(os.path.join(OUT_DIR, "simscale_fem_stress_fine.png"), width=Inches(5.5))
            
            p2 = p.insert_paragraph_before()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run()
            run2.add_picture(os.path.join(OUT_DIR, "simscale_fem_displacement.png"), width=Inches(5.5))
            
        elif "(1) pressure contour" in p.text:
            p.text = "SimScale Computational Fluid Dynamics (CFD) Results — Surface Pressure Coefficient & Velocity Streamlines:"
            p.runs[0].font.bold = True
            
            # Add images
            p1 = p.insert_paragraph_before()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p1.add_run()
            run1.add_picture(os.path.join(OUT_DIR, "simscale_cfd_pressure.png"), width=Inches(5.5))
            
            p2 = p.insert_paragraph_before()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run()
            run2.add_picture(os.path.join(OUT_DIR, "simscale_cfd_streamlines.png"), width=Inches(5.5))

doc.save(DOC_PATH)
print(f"Successfully updated and saved {DOC_PATH}!")
